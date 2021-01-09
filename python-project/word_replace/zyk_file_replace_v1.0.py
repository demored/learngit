# -*-coding:utf-8-*-

'''
安装win32com
python -m pip install pypiwin32
安装docx
pip install python-docx
'''
import os
import sys
import docx
from win32com import client as wc
import fileinput
import openpyxl
import re
import traceback
import shutil
import xlrd
import xlwt

# 定义文件目录为当前目录
work_dir = os.path.abspath('.')
replace_dir = work_dir+"\\替换目录\\"
result_dir = work_dir+"\\替换结果\\"

try:
    word = wc.Dispatch('kwps.Application')
except:
    word = wc.Dispatch('word.Application')

#获取目录下所有doc文件
def get_dir_doc_files(filepath):
    # doc文件列表
    doc_files = []
    # #读取文件夹中文件列表
    for file in os.listdir(filepath):
        if file.endswith(".doc") and not file.startswith('~$'):
            doc_files.append(file)
    return doc_files

#获取目录中所有docx文件
def get_dir_docx_files(filepath):
    # 获取目录下所有doc文件
    # docx文件列表
    docx_files = []
    # #读取文件夹中文件列表
    for file in os.listdir(filepath):
        if file.endswith(".docx") and not file.startswith('~$'):
            docx_files.append(file)
    return docx_files

#获取目录中所有的excel文件
def get_dir_excel_files(filepath):
    # 获取目录下所有doc文件
    # docx文件列表
    excel_files = []
    # #读取文件夹中文件列表
    for file in os.listdir(filepath):
        if file.endswith(".xls") and not file.startswith('~$'):
            excel_files.append(file)
        elif file.endswith(".xlsx") and not file.startswith('~$'):
            excel_files.append(file)

    return excel_files

#将old_path中的docx，xls,xlsx文件复制到new_path中
def mv_file_result_dir(old_path, new_path):
    for file in os.listdir(old_path):
        # if file.endswith(".xls") and not file.startswith('~$'):
        #     shutil.copyfile(old_path+file, new_path+file)
        # elif file.endswith(".xlsx") and not file.startswith('~$'):
        #     shutil.copyfile(old_path + file, new_path + file)
        if file.endswith(".docx") and not file.startswith('~$'):
            shutil.copyfile(old_path + file, new_path + file)


# #将doc替换成docx
def doSaveAas(files_arr):
    for i in files_arr:
        temp = replace_dir+'{}'.format(i)
        doc = word.Documents.Open(temp)  #目标路径下的文件
        rename = os.path.splitext(i)
        doc.SaveAs(result_dir + rename[0] + '.docx', 12)  # 12表示docx格式
        doc.Close()
    word.Quit()

#文字替换
def info_update(doc,old_info, new_info):
    '''此函数用于批量替换合同中需要替换的信息
    doc:文件
    old_info和new_info：原文字和需要替换的新文字
    '''
    changeCells_word_content = 0
    changeCells_word_table = 0

    #替换页眉
    # for yemei_i in doc.sections:
    #     head = doc.sections[yemei_i].header
    #     for para in doc.paragraphs:
    #         for run in para.runs:

    #读取段落中的所有run，找到需替换的信息进行替换
    for para in doc.paragraphs: #
        for run in para.runs:
            #统计字符串出现次数
            print(run.text)
            changeCells_word_content += run.text.count(old_info)
            # pattern = re.compile(old_info)
            run.text = run.text.replace(old_info, new_info, 1) #替换信息
            # run.text = re.sub(old_info,new_info,run.text)

    #读取表格中的所有单元格，找到需替换的信息进行替换
    # for table in doc.tables:
    #     for row in table.rows:
    #         for cell in row.cells:
    #             changeCells_word_table += cell.text.count(old_info)
    #             cell.text = cell.text.replace(old_info, new_info) #替换信息

    print(changeCells_word_content)
    # print(changeCells_word_table)
    return changeCells_word_content + changeCells_word_table

if __name__ == '__main__':
    print("证优客咨询师文件替换工具V1.0")
    # 创建需要替换的目录
    if not os.path.exists(replace_dir):
        os.makedirs(replace_dir)

    # 创建替换后的目录
    if not os.path.exists(result_dir):
        os.makedirs(result_dir)

    # 如果存在doc文件则将doc转换成docx
    doc_files = get_dir_doc_files(replace_dir)
    if len(doc_files) > 0:
        doSaveAas(doc_files)
    #将replace中的不需要转换的文件复制到result目录
    mv_file_result_dir(replace_dir, result_dir)

    #获取需要替换的文件
    replace_files = os.listdir(result_dir)
    ##读取配置文件
    config_text = work_dir +"\\" +"config.txt"

    if not os.path.exists(config_text):
        print("配置文件不存在")
        sys.exit()

    for i in replace_files:
        #加上绝对路径
        file = result_dir + i
        print("==============开始"+file+"关键字替换==============")

        #判断是docx还是excel
        # if i.endswith(".docx") and not i.startswith('~$'):
        doc = docx.Document(file)

        f = open(config_text, encoding='utf-8')
        while True:
            line = f.readline()
            if line:
                split_list = line.split('|')
                old_word = split_list[0].replace('\r','').replace('\n','').replace('\t','')
                new_word = split_list[1].replace('\r','').replace('\n','').replace('\t','')
                nums = info_update(doc, old_word, new_word)
                print("【%s】替换为【%s】共替换：【%d】次" % (old_word, new_word,nums))
            else:
                break
        f.close()
        doc.save("{0}{1}".format(result_dir,file.split("\\")[-1]))
    os.system("pause")