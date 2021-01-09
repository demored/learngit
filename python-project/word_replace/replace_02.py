from docx import Document
import os
def change_text(old_text, new_text):
    # 更换段中的文字
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.text = run.text.replace(old_text, new_text)

    # 更换表格中的文字
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.text = run.text.replace(old_text, new_text)

work_dir = os.path.abspath('.')
file_name = work_dir+"\\"+"[CM-IM-004-V1.0]信息安全目标.docx"
document = Document(file_name)
change_text('蝉鸣科技（西安）有限公司', '证优客')

save_name = work_dir+"\\" +"demo.docx"
document.save(save_name)